import os
from pathlib import Path
from groq import Groq
from dotenv import load_dotenv
load_dotenv()
api_key=os.getenv("GROQ_API_KEY")
if not api_key:
    raise ValueError("API Key not available in .env")
client=Groq(api_key=api_key)
model="llama-3.3-70b-versatile"

#step1:JD
job_description='''
Description
Job Brief
‍

We are seeking for AI (Artificial Intelligence) engineers to work with us! In this role, you would create and use artificial intelligence (AI) within projects. You should be well-versed in algorithms that can evolve and grow into other processes.

‍

Additionally, you would be expected to know about data science fundamentals and programming languages, such as Python and C++, to develop software solutions that satisfy project specifications and match the demand for artificial intelligence applications. Your tasks will also include the creation of new models and algorithms to address issues in collaboration with other team members of software engineers, data scientists, and other professionals.

‍

‍

Roles and Responsibilities
‍

Develop AI strategies through research and development methods.
Work with cross-functional teams to highlight areas where using AI may promote corporate success.
Build an effective prototype.
Perform tests and troubleshoot issues before the final delivery.
Keep up with the most recent advancements in artificial intelligence that affect the corporate sector.
Manage and direct processes and R&D (research and development) to meet the needs of our AI strategy.
Identify the problems your business and clients are experiencing and how incorporating AI skills might assist in finding solutions.
When analyzing and explaining AI and machine learning (ML) solutions, set and maintain excellent ethical standards.
Advising corporate leaders and C-suite executives on a variety of technical, strategy, and policy challenges relating to AI
Work on functional design, process design (scenario design and flow mapping), prototyping, training, and creating support processes.
Articulate and document the solutions architecture and lessons learned for each exploration and accelerated incubation.
Serve as liaison between stakeholders and project teams, providing input and helping team members to make the required adjustments to the product's performance or presentation.
‍

‍

Requirements and Skills
‍

B.Tech (Bachelor of Technology) degrees in computer science engineering, artificial intelligence, or a related field are necessary.
Knowledge of AI's most widely used programming languages, including Python, Java, C++, and R.
Deep understanding of statistical and algorithmic models as well as fundamental mathematical concepts like probability and linear algebra
Familiarity with building effective programs that can quickly analyze massive data streams and working with large data sets
Employment of well-known AI/ML frameworks, familiarity with deep learning and machine learning techniques
Proven experience in applying AI to practical and all-inclusive technology solutions
Hands-on experience in deep learning, Python, Tensorflow, and machine learning
A better understanding of fundamental algorithms, functional design ideas, and object-oriented programming principle
Experience in designing NoSQL databases, designing RDBMSs, and developing REST APIs
Knowledge of TensorFlow, PyTorch, and Keras, three popular AI programming frameworks.
Expertise in data analysis and visualization programs like Matplotlib, Power BI, and Tableau.
Abilities work on computer vision (CV) and natural language processing (NLP).
Knowledge of deep learning frameworks such as recurrent and convolutional neural networks (RNNs).
Good organizational, analytical, and communication skills are needed.
Ability to function both individually and as a team
‍

'''
#step2: JobD class,schema,system prompt
from pydantic import BaseModel,Field
class JobD(BaseModel):
    role:str
    required_skills:list[str]
    preferred_skills:list[str]
    minimum_experience:float|None
    education_requirements:list[str]
    responsibilities:list[str]
jobd_schema=JobD.model_json_schema()

system_prompt=f'''
You are an expert HR Assistant
Your job is to analyze job description and extract structured information from it
Return only valid json matching this schema {jobd_schema} 
Note:
Do not return the schema itself
Do not return fields like "properties","titles",etc which are not in the schema
Fill the schema with actual information extracted from job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list
do not invent information 
'''
user_prompt=f'''
Analyze the following job description:
{job_description}
'''
message_system={
    "role":"system",
    "content":system_prompt
}
message_user={
    "role":"user",
    "content":user_prompt
}
response_format={
    "type":"json_object"
}
messages=[message_system,message_user]
response=client.chat.completions.create(model=model,messages=messages,response_format=response_format)
answer=response.choices[0].message.content

#step3 JD extracted json into reading its parameters
raw_json=answer
print(raw_json)
import json
job_data=json.loads(raw_json)
job=JobD(**job_data)
print(f"Minimum experience required:{job.minimum_experience}")
print(f"job education req:{job.education_requirements}")

#step4: Resume class, schema, parsing 
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []
class Resume(BaseModel):
    name:str |None=None
    email: str | None = None
    phone: str | None = None
    total_experience_years: float | None = None
    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []
resume_schema=Resume.model_json_schema()
class MatchResult(BaseModel):
    score: float
    details: dict

#step5:Final score function using llm 
def final_score(job, resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are a strict HR recruiter. Compare the candidate's resume with the job description.
    Calculate a precise match score from 0 to 100 based strictly on:
    - How many required skills are present in the resume.
    - Whether the education requirements are met.
    - Whether the experience requirements are met.
    
    Do NOT give a generic score. If critical skills are missing, the score must drop significantly (e.g., below 50).

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}
    
    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    
    Return JSON matching this schema: {match_schema}
    Give me:
    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met (true/false)
    5. Overall match percentage from 0 to 100 (calculated strictly based on gaps)
    6. A short final verdict
    """
    message = {"role": "user", "content": prompt}
    response_format = {"type": "json_object"}
    response = client.chat.completions.create(
        model=model, messages=[message], response_format=response_format
    )
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)

#step6:Parse resume
def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume

#step7:reading resume
from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
    reader=PdfReader(file_path)
    text=""
    for page in reader.pages:
        page_text=page.extract_text()
        if page_text:
            text+=page_text+"\n"
    return text
def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text
def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None

#Main function
import time
resume_folder = Path("resumes")
all_results=[]
for file_path in resume_folder.iterdir():
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume=parse_resume(resume_text) # llm call1
    time.sleep(5)
    result = final_score(job, parsed_resume) #llm caLL2
    time.sleep(5)
    print("Score:", result.score)
    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    })
all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)
top_2 = all_results[:2]
worst_2 = all_results[-2:]


print("TOP 2 CANDIDATES")
for candidate in top_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])

